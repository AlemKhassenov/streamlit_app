import streamlit as st
import pandas as pd
import requests
import io
import os
from docx import Document

# Загружаем API-ключ из переменной окружения
API_KEY = os.getenv("OPENAI_API_KEY")

# Проверка API-ключа
if not API_KEY:
    st.error("Ошибка: API-ключ не найден. Добавьте его в переменные среды или GitHub Secrets.")
    st.stop()

# Функция для загрузки данных из Google Sheets
def load_data_from_google_sheets(student_name, student_class):
    try:
        sheet_url = "https://docs.google.com/spreadsheets/d/1BeXPi5LRSIj0xDjGZjey0VfBU08mnjJm/export?format=xlsx"
        xls = pd.ExcelFile(sheet_url)
        student_data_list = []
        
        # Ищем сначала в указанной пользователем вкладке
        if student_class in xls.sheet_names:
            sheet_names = [student_class] + [s for s in xls.sheet_names if s != student_class]
        else:
            sheet_names = xls.sheet_names
        
        for sheet_name in sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            df.columns = df.columns.str.strip()
            
            # Определяем правильное название столбца с ФИО
            fio_column = next((col for col in df.columns if "ФИО" in col), None)
            
            if fio_column:
                df[fio_column] = df[fio_column].fillna('').astype(str).str.strip()
                
                student_data = df[df[fio_column] == student_name.strip()]
                
                if not student_data.empty:
                    # Формируем словарь, исключая сам столбец ФИО
                    student_info = student_data.drop(columns=[fio_column]).to_dict(orient='records')[0]
                    student_data_list.append(student_info)
                    break  # Если нашли данные, выходим из цикла
        
        if not student_data_list:
            return None
        
        return student_data_list[0]  # Возвращаем первую найденную строку как словарь без ФИО
    except Exception as e:
        st.error(f"Ошибка при загрузке данных из Google Sheets: {e}")
        return None

# Функция для отправки данных через API ChatGPT
def send_data_to_api(student_data, prompt):
    api_url = "https://api.openai.com/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    final_prompt = f"{prompt}\n\nДанные ученика:\n{student_data}"

    payload = {
        "model": "gpt-4",
        "messages": [{"role": "user", "content": final_prompt}],
        "max_tokens": 1000,
        "temperature": 0.7
    }

    response = requests.post(api_url, json=payload, headers=headers)
    
    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        return f"Ошибка API: {response.status_code}, {response.text}"

# Функция для создания DOCX-файла
def create_docx(student_name, student_class, subject, grade, gpt_response):
    doc = Document()
    
    doc.add_heading("Индивидуальный план развития ученика", level=1)

    doc.add_paragraph(f"ФИО ученика: {student_name}")
    doc.add_paragraph(f"Класс: {student_class}")
    doc.add_paragraph(f"Предмет: {subject}")
    doc.add_paragraph(f"Текущая оценка: {grade}")

    doc.add_heading("Анализ и рекомендации", level=2)
    doc.add_paragraph(gpt_response)

    # Сохранение файла в буфер
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Функция для установки фонового изображения
def set_background(image_url):
    page_bg = f"""
    <style>
    .stApp {{
        background: url("{image_url}") no-repeat center center fixed;
        background-size: cover;
    }}
    </style>
    """
    st.markdown(page_bg, unsafe_allow_html=True)

# Укажите ссылку на фоновое изображение
background_image_url = "https://images.pexels.com/photos/30388784/pexels-photo-30388784.jpeg?auto=compress&cs=tinysrgb&w=1200&lazy=load"
set_background(background_image_url)

# Словарь соответствий классов и учеников
students_by_class = {
    "5А (Final1)": ["Абешев Арман", "Аблазим Нұрамир", "Амиртай Абдулла", "Аубакир Амандос", "Әмен Айни", "Бигазина Дарина", "Даулет Еркежан", "Жұматай Аянат", "Заит Мирас", "Кумайхан Адия", "Куренбекова Ажар", "Қайырбай Манас", "Мәди Аяжан", "Мұхтар Нұрали", "Наурызбай Бекарыс", "Толегенова Фатима", "Тығынбай Бексултан", "Фаридқызы Айсұлу"],
    "5 E,F,D(Final1)": ["Айдарова Аяна", "Ақтан Мадина", "Брюханова Кира", "Гусев Матвей", "Джамбулова Жасмин", "Дидар Дарина", "Жанболат Гүлжан", "Закарин Ильяс", "Курмантаева Амина", "Қожахмет Қасым", "Марат Райана", "Мухамеджанов Тимур", "Мұхарап Айсұлтан", "Райхан Аскар", "Рамазанов Абзал", "Рахматуллин Ибрагим", "Рожков Игнатий", "Тлеулин Турар", "Төлеухан Жанибек", "Уйсумбаева Ханифа"],
    "6 A (Final1)": ["Абдсаттар Алишер", "Аблазим Жанасыл", "Абусаид Алан", "Аманкелдиев Мансур", "Ахметова Айдана", "Ахметова Жәмиля", "Берікова Қымбат", "Болат Мирас", "Жагипаров Алимхан", "Жұматай Іңкар", "Мурзатаев Асылхан", "Нуракишев Шыңғыс", "Ордабаева Аиша", "Сарыбай Инара", "Саурбаева Айша", "Умирзакова Алина"],
    "7a Final 1": ["Ақкісі, Ансар", "Аяпберген, Исмаил", "Бегешова, Баян", "Бейбітұлы, Сағди", "Берикбол, Рамина", "Богданова, Айя", "Даниярқызы, Адия", "Зейнешев, Санжар", "Зиябек, Аиша", "Каирбек, Жасмин", "Каирбек, Нұрмұхаммед", "Келімбет, Даниал", "Нұрқиянов, Аңсар", "Рахметулла, Айсана", "Сериков, Аслан", "Серикова, Шадия", "Скендер, Айсара", "Тлеухан, Айнаш", "Тюлеубай, Еркежан"]
}

# Заголовок
st.title("Индивидуальный план развития ученика Quantum STEM School")

# Выбор класса
student_class = st.selectbox("Класс", list(students_by_class.keys()))

# После выбора класса отображаем учеников
if student_class:
    student_name = st.selectbox("ФИО ученика", students_by_class[student_class])
    
    with st.form("student_form"):
        subject = st.selectbox("Предмет", ["Математика", "Физика", "Химия", "Биология", "Английский язык"])
        grade = st.number_input("Текущая оценка из EduPage, округленная до целых", min_value=0, max_value=100, step=1, value=0)
        submit_button = st.form_submit_button("Сформировать ПИР")

    if submit_button:
        st.write("Загрузка данных из Google Sheets...")
        student_data = load_data_from_google_sheets(student_name, student_class)
        
        if student_data is None:
            st.error("Ошибка: Данные ученика не найдены! Проверьте ФИО и структуру таблицы.")
        else:
            st.write("Данные найдены, отправляются на анализ...")
            st.write(student_data)

if submit_button:
    st.write("Загрузка данных из Google Sheets...")
    student_data = load_data_from_google_sheets(name, student_class)
    
    if student_data is None:
        st.error("Ошибка: Данные ученика не найдены! Проверьте ФИО и структуру таблицы.")
    else:
        st.write("Данные найдены, отправляются на анализ...")
        
        prompt = """
Учащийся проходил обучение по различным темам. В столбцах приведены ожидания от обучения (цели), а в последней строке указано:
- 1 = ученик достиг цели.
- 0 = ученик не достиг цели.

Проанализируй, какие темы усвоены хорошо, а какие требуют доработки.  
Определи, какие пробелы в знаниях есть у ученика, и предложи **рекомендации**.  
Для каждой недостигнутой цели (0) предложи **ресурсы для изучения**: 
- **Ссылки на видео (YouTube, Coursera, Khan Academy и т. д.)**  
- **Книги / статьи / курсы**  

Данные для анализа:  
{student_data}
"""
        api_response = send_data_to_api(student_data, prompt)
        
        if "Ошибка" in api_response:
            st.error(api_response)
        else:
            doc_buffer = create_docx(name, student_class, subject, grade, api_response)
            st.download_button("📄 Скачать отчет (DOCX)", data=doc_buffer, file_name=f"ПИР_{name}_{grade}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.success("Отчет сформирован! Вы можете скачать его.")
